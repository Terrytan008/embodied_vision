#!/usr/bin/env python3
"""
报销单填写脚本
根据发票数据和模板填写报销单
"""

import sys
import json
import argparse
from pathlib import Path

def fill_excel_form(data, template_path, output_path):
    """填写 Excel 报销单"""
    try:
        import openpyxl
    except ImportError:
        print("错误: 需要安装 openpyxl pip install openpyxl", file=sys.stderr)
        sys.exit(1)
    
    wb = openpyxl.load_workbook(template_path)
    ws = wb.active
    
    # 字段映射（根据实际模板调整）
    field_map = {
        'A1': 'applicant',
        'B1': 'department',
        'C1': 'reimbursement_date',
    }
    
    # 填写基本信息
    for cell, field in field_map.items():
        if field in data:
            ws[cell] = data[field]
    
    # 填写报销明细（从第3行开始）
    if 'expenses' in data:
        for i, expense in enumerate(data['expenses']):
            row = i + 3
            ws[f'A{row}'] = expense.get('date', '')
            ws[f'B{row}'] = expense.get('type', '')
            ws[f'C{row}'] = expense.get('description', '')
            ws[f'D{row}'] = expense.get('amount', 0)
            ws[f'E{row}'] = expense.get('invoice_number', '')
    
    # 填写合计
    if 'total_amount' in data:
        ws['D10'] = data['total_amount']  # 假设合计在 D10
    
    wb.save(output_path)
    return True

def fill_word_form(data, template_path, output_path):
    """填写 Word 报销单"""
    try:
        from docx import Document
    except ImportError:
        print("错误: 需要安装 python-docx pip install python-docx", file=sys.stderr)
        sys.exit(1)
    
    doc = Document(template_path)
    
    # 替换占位符 {{field_name}}
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f'{{{{{key}}}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))
    
    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if f'{{{{{key}}}}}' in cell.text:
                        cell.text = cell.text.replace(f'{{{{{key}}}}}', str(value))
    
    doc.save(output_path)
    return True

def main():
    parser = argparse.ArgumentParser(description='报销单填写')
    parser.add_argument('--data', '-d', required=True, help='报销数据 JSON 文件')
    parser.add_argument('--template', '-t', required=True, help='报销单模板文件')
    parser.add_argument('--output', '-o', required=True, help='输出文件路径')
    
    args = parser.parse_args()
    
    # 读取数据
    with open(args.data, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    template_path = Path(args.template)
    output_path = Path(args.output)
    
    if not template_path.exists():
        print(f"错误: 模板文件不存在 {args.template}", file=sys.stderr)
        sys.exit(1)
    
    # 根据模板类型选择填写方式
    if template_path.suffix == '.xlsx':
        success = fill_excel_form(data, args.template, args.output)
    elif template_path.suffix == '.docx':
        success = fill_word_form(data, args.template, args.output)
    else:
        print(f"错误: 不支持的模板格式 {template_path.suffix}", file=sys.stderr)
        sys.exit(1)
    
    if success:
        print(f"报销单已生成: {args.output}")

if __name__ == '__main__':
    main()
